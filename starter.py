import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import docx

import sys
print(sys.executable)
# 3-phase

pi = np.pi 

x = np.linspace(0,5*pi,100)
y1 = np.exp(-0.3338*x)*np.sin(x + pi/3)
y2 = np.exp(-0.4325*x)*np.sin(x + 2*pi/3)
y3 = np.exp(-0.9123*x)*np.sin(x + 3*pi/3)

plt.figure(1)
plt.plot(x,y1,'b-')
plt.plot(x,y2,'r-')
plt.plot(x,y3,'k-')
plt.grid(True)
plt.show()


df = pd.DataFrame(data = {'y1': y1,
                          'y2': y2,
                          'y3': y3,
                          
                          })

print(df)
df.to_excel('3-phase.xlsx', index=False)

# ...existing code...

# Read Excel data
df = pd.read_excel('3-phase.xlsx')

# Create a Word document
doc = docx.Document()
doc.add_heading('3-Phase Data', 0)

# Add table
table = doc.add_table(rows=1, cols=len(df.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(df.columns):
    hdr_cells[i].text = col

for idx, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

doc.save('3-phase.docx')
